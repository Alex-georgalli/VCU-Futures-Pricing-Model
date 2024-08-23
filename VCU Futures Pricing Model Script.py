# In[1] Dependencies
import numpy as np
import matplotlib.pyplot as plt
from scipy.optimize import minimize
from scipy.stats import gaussian_kde
from scipy.stats import kstest
import seaborn as sns
import pandas as pd
from docx import Document  # pip install python-docx (if not already installed) (used for results report)

# In[2] Parameters
security = 'YBBZ4 Comdty'
lookback = 252 #how long the total days should be (pre in and out split)
in_sample_percentage = 0.80 #(% of total days that should be in sample)
Nsim_fixed = 10000 #(number of simulations per model when running the minimizatoin)
Nruns = 15 #(number of times to run each model minimization in order to calcualte the average)
seed_fixed = None

# In[3] Data Importing
data = pd.read_excel(r"C:") #Enter the file path for the Data Master YBBZ4 xlxs file
data.dropna()
credit_prices = data.set_index('Day')[security]
credit_prices = credit_prices.dropna()
credit_prices = credit_prices.iloc[-lookback:]
first_day = credit_prices.index[0]
total_days = len(credit_prices)
in_sample_length = int(in_sample_percentage * total_days)
out_of_sample_length = total_days - in_sample_length
in_sample_data = credit_prices.iloc[:in_sample_length]
out_of_sample_data = credit_prices.iloc[in_sample_length:]
real_log_returns = np.diff(np.log(credit_prices.values))

# In[4] Economic Evolutions
def relative_sector_evolution_params_only(A_c0, A_d0, eta_prod, alpha, epsilon,
                              Theta_good, c, sigma_carbon, Income, q_t, num_periods, plot=False):
    
    gamma = (1 - alpha) * (epsilon - 1)
    theta = Theta_good ** epsilon
    denom = c * sigma_carbon * Income

    # Arrays to store results
    Rel_prod_array = np.zeros(num_periods)  # time t array
    Rel_profit_array = np.zeros(num_periods)  # time t array
    Rel_rnd_array = np.zeros(num_periods)  # time t array
    Carbon_return_array = np.zeros(num_periods)  # time t array

    Rel_prod_array_lag_1 = np.zeros(num_periods)  # time t-1 array
    Rel_prod_array_lag_1[0] = A_c0 / A_d0  # Initial value at time 0

    # Model Loop
    A_ct = A_c0
    A_dt = A_d0

    # Evolution Paths
    for t in range(0, num_periods):  # Start from time 1
        A_ct_tilda_approx = ((A_ct / A_dt) ** - 1) / (1 + ((A_ct / A_dt) ** - 1))
        A_dt_tilda_approx = ((A_dt / A_ct) ** - 1) / (1 + ((A_dt / A_ct) ** - 1))

        A_ct_approx_new = A_ct * (1 + eta_prod * A_ct_tilda_approx)
        A_dt_approx_new = A_dt * (1 + eta_prod * A_dt_tilda_approx)

        g = (1 + eta_prod * A_ct_tilda_approx) / (1 + eta_prod * A_dt_tilda_approx)
        g_var = (g * (A_ct / A_dt)**gamma)

        # Store the values for the evolution of relative sector productivities
        Rel_prod_array[t] = A_ct_approx_new / A_dt_approx_new

        if t > 0:  
            Rel_prod_array_lag_1[t] = A_ct / A_dt
        
        Carbon_return_array[t] = 1/denom * ((g_var*(theta+1)+1+theta*(g_var**2)) /\
                                            (-q_t*g_var*(theta-1)+1+theta*g_var))**2 - 1

        # Update for next iteration
        A_ct = A_ct_approx_new
        A_dt = A_dt_approx_new
        
    #Plotting 2x2 subplot
    if plot:
        plt.figure(figsize=(10, 8))
        plt.suptitle('$A_{{c0}}$ = {:.2f}, $A_{{d0}}$ = {:.2f}, $η$ = {:.2f}, $ε$ = {:.2f}, $α$ = {:.2f}, $θ$ = {:.2f}, $c$ = {:.2f}, $σ_{{c}}$ = {:.2f}, $I_{{t}}$ = {:.2f}, $q_{{t}}$ = {:.2f}'.format(A_c0, A_d0, eta_prod, epsilon, alpha, Theta_good, c, sigma_carbon, Income, q_t), fontsize=14)

        # Plotting Relative Productivity
        plt.subplot(2, 2, 1)
        plt.plot(range(num_periods), Rel_prod_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$A_{ct}/A_{dt}$')
        plt.title('Evolution of Relative Sector Productivities over Time')
        plt.grid(True)

        # Plotting Relative Profit over time
        plt.subplot(2, 2, 2)
        plt.plot(range(num_periods), Rel_profit_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$π_{ct}/π_{dt}$')
        plt.title('Evolution of Relative Sector Proftis over Time')
        plt.grid(True)

        # Plotting Carbon Return over time
        plt.subplot(2, 2, 3)
        plt.plot(range(num_periods), Carbon_return_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$μ_{t}$')
        plt.title('Expected Carbon Credit Return over Time')
        plt.grid(True)

        # Plotting Relative R&D over time
        plt.subplot(2, 2, 4)
        plt.plot(range(num_periods), Rel_rnd_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$R_{ct}/R_{dt}$')
        plt.title('Evolution of Relative Sector R&D Expenditure over Time')
        plt.grid(True)

        plt.tight_layout()
        plt.show()
    
    return Carbon_return_array

def relative_sector_evolution_with_patents(patent_data, eta_prod, alpha, epsilon,
                              Theta_good, c, sigma_carbon, Income, q_t, num_periods, plot=False):
  
    gamma = (1 - alpha) * (epsilon - 1)
    theta = Theta_good ** epsilon
    denom = c * sigma_carbon * Income

    # Arrays to store results
    Rel_prod_array = np.zeros(num_periods)  # time t array
    Rel_profit_array = np.zeros(num_periods)  # time t array
    Rel_rnd_array = np.zeros(num_periods)  # time t array
    Carbon_return_array = np.zeros(num_periods)  # time t array

    Rel_prod_array_lag_1 = np.zeros(num_periods)  # time t-1 array

    # Initialize A_c0 and A_d0
    A_c0 = patent_data['Clean'].iloc[0]
    A_d0 = patent_data['Dirty'].iloc[0]
    
    Rel_prod_array_lag_1[0] = A_c0 / A_d0  # Initial value at time 0

    # Model Loop
    for t in range(num_periods):  # Start from time 1
        # Update A_c0 and A_d0 based on patent data
        A_ct = patent_data['Clean'].iloc[t]
        A_dt = patent_data['Dirty'].iloc[t]
        
        # Calculate A_ct/A_dt ratio
        Rel_prod_array_lag_1[t] = A_ct / A_dt
        
        A_ct_tilda_approx = ((A_ct / A_dt) ** - 1) / (1 + ((A_ct / A_dt) ** - 1))
        A_dt_tilda_approx = ((A_dt / A_ct) ** - 1) / (1 + ((A_dt / A_ct) ** - 1))

        A_ct_approx_new = A_ct * (1 + eta_prod * A_ct_tilda_approx)
        A_dt_approx_new = A_dt * (1 + eta_prod * A_dt_tilda_approx)

        g = (1 + eta_prod * A_ct_tilda_approx) / (1 + eta_prod * A_dt_tilda_approx)
        g_var = (g * (A_ct / A_dt)**gamma)

        # Store the values for the evolution of relative sector productivities
        Rel_prod_array[t] = A_ct_approx_new / A_dt_approx_new

        if t > 0:  
            Rel_prod_array_lag_1[t] = A_ct / A_dt
        
        Carbon_return_array[t] = 1/denom * ((g_var*(theta+1)+1+theta*(g_var**2)) /\
                                            (-q_t*g_var*(theta-1)+1+theta*g_var))**2 - 1

        # Update for next iteration
        A_ct = A_ct_approx_new
        A_dt = A_dt_approx_new

    if plot:
        plt.figure(figsize=(10, 8))
        plt.suptitle('$η$ = {:.2f}, $ε$ = {:.2f}, $α$ = {:.2f}, $θ$ = {:.2f}, $c$ = {:.2f}, $σ_{{c}}$ = {:.2f}, $I_{{t}}$ = {:.2f}, $q_{{t}}$ = {:.2f}'.format(eta_prod, epsilon, alpha, Theta_good, c, sigma_carbon, Income, q_t), fontsize=14)

        # Plotting Relative Productivity
        plt.subplot(2, 2, 1)
        plt.plot(range(num_periods), Rel_prod_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$A_{ct}/A_{dt}$')
        plt.title('Evolution of Relative Sector Productivities over Time')
        plt.grid(True)

        # Plotting Relative Profit over time
        plt.subplot(2, 2, 2)
        plt.plot(range(num_periods), Rel_profit_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$π_{ct}/π_{dt}$')
        plt.title('Evolution of Relative Sector Proftis over Time')
        plt.grid(True)

        # Plotting Carbon Return over time
        plt.subplot(2, 2, 3)
        plt.plot(range(num_periods), Carbon_return_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$μ_{t}$')
        plt.title('Expected Carbon Credit Return over Time')
        plt.grid(True)

        # Plotting Relative R&D over time
        plt.subplot(2, 2, 4)
        plt.plot(range(num_periods), Rel_rnd_array)
        plt.xlabel('Period (t)')
        plt.ylabel('$R_{ct}/R_{dt}$')
        plt.title('Evolution of Relative Sector R&D Expenditure over Time')
        plt.grid(True)

        plt.tight_layout()
        plt.show()
    
    return Carbon_return_array

# In[5] Diffusion Processes
def economic_jump_diffusion(S, T, mu_array, sigma, Lambda, a, b, Nsteps, Nsim, seed, plot=False):
    
    # Set random seed
    np.random.seed(seed)

    # Calculate the length of the time step
    Delta_t = T/Nsteps

    simulated_paths = np.zeros([Nsim, Nsteps+1])

    # Replace the first column of the array with the vector of initial price S
    simulated_paths[:,0] = S

    Z_1 = np.random.normal(size=[Nsim, Nsteps])
    Z_2 = np.random.normal(size=[Nsim, Nsteps])
    Poisson = np.random.poisson(Lambda*Delta_t, [Nsim, Nsteps])

    # Populate the matrix with Nsim randomly generated paths of length Nsteps
    for i in range(Nsteps):
        mu = mu_array[i]  # Use the corresponding mu value from the array
        
        simulated_paths[:,i+1] = simulated_paths[:,i]*np.exp((mu
                               - sigma**2/2)*Delta_t + sigma*np.sqrt(Delta_t) \
                               * Z_1[:,i] + a*Poisson[:,i] \
                               + np.sqrt(b**2) * np.sqrt(Poisson[:,i]) \
                               * Z_2[:,i])

    # Calculate the average path from all simulated paths
    mean_path = np.mean(simulated_paths, axis=0)

    if plot:
        # Choose palette, figure size, and define figure axes
        sns.set(palette='rocket_r')
        plt.figure(figsize=(10,8))
        
        ax = plt.axes()
        t = np.linspace(0, T, Nsteps+1) * Nsteps
        jump_diffusion = ax.plot(t, simulated_paths.transpose());
        plt.setp(jump_diffusion, linewidth=1)
        plt.plot(t, mean_path, color='cyan', linewidth=2, label='Mean Path');
        ax.set(title="Monte Carlo simulated price paths\n$S_0$ = {}, $sigma$ = {}, $a$ = {}, $b$ = {}, $lambda$ = {}, Nsim = {}"\
               .format(S, sigma, a, b, Lambda, Nsim), \
               xlabel='Time (days)', ylabel=f"{security}")
        plt.legend()
        plt.show()
    
    return mean_path

def pure_jump_diffusion(S, T, mu, sigma, Lambda, a, b, Nsteps, Nsim, seed, plot=False):
    
    # Set random seed
    np.random.seed(seed)

    # Calculate the length of the time step
    Delta_t = T/Nsteps

    simulated_paths = np.zeros([Nsim, Nsteps+1])

    # Replace the first column of the array with the vector of initial price S
    simulated_paths[:,0] = S

    Z_1 = np.random.normal(size=[Nsim, Nsteps])
    Z_2 = np.random.normal(size=[Nsim, Nsteps])
    Poisson = np.random.poisson(Lambda*Delta_t, [Nsim, Nsteps])

    # Populate the matrix with Nsim randomly generated paths of length Nsteps
    for i in range(Nsteps):
        simulated_paths[:,i+1] = simulated_paths[:,i]*np.exp((mu
                               - sigma**2/2)*Delta_t + sigma*np.sqrt(Delta_t) \
                               * Z_1[:,i] + a*Poisson[:,i] \
                               + np.sqrt(b**2) * np.sqrt(Poisson[:,i]) \
                               * Z_2[:,i])

    mean_path = np.mean(simulated_paths, axis=0)
    
    if plot:
        sns.set(palette='rocket_r')
        plt.figure(figsize=(10,8))
        
        ax = plt.axes()
        t = np.linspace(0, T, Nsteps+1) * Nsteps
        jump_diffusion = ax.plot(t, simulated_paths.transpose());
        plt.setp(jump_diffusion, linewidth=1)
        plt.plot(t, mean_path, color='cyan', linewidth=2, label='Mean Path');
        ax.set(title="Monte Carlo simulated price paths\n$S_0$ = {}, $sigma$ = {}, $a$ = {}, $b$ = {}, $lambda$ = {}, Nsim = {}"\
               .format(S, sigma, a, b, Lambda, Nsim), \
               xlabel='Time (days)', ylabel=f"{security}")
        plt.legend()
        plt.show()
    
    return mean_path

def economic_drift_diffusion(S, T, mu_array, sigma, Nsteps, Nsim, seed, plot=False):
    
    # Set random seed
    np.random.seed(seed)

    # Calculate the length of the time step
    Delta_t = T/Nsteps

    simulated_paths = np.zeros([Nsim, Nsteps+1])

    # Replace the first column of the array with the vector of initial price S
    simulated_paths[:,0] = S

    Z_1 = np.random.normal(size=[Nsim, Nsteps])

    # Populate the matrix with Nsim randomly generated paths of length Nsteps
    for i in range(Nsteps):
        mu = mu_array[i]  # Use the corresponding mu value from the array
        
        simulated_paths[:,i+1] = simulated_paths[:,i]*np.exp((mu - sigma**2/2)*Delta_t +\
                                                             sigma*np.sqrt(Delta_t) * Z_1[:,i])

    # Calculate the average path from all simulated paths
    mean_path = np.mean(simulated_paths, axis=0)

    if plot:
        # Choose palette, figure size, and define figure axes
        sns.set(palette='rocket_r')
        plt.figure(figsize=(10,8))
        
        ax = plt.axes()
        t = np.linspace(0, T, Nsteps+1) * Nsteps
        jump_diffusion = ax.plot(t, simulated_paths.transpose());
        plt.setp(jump_diffusion, linewidth=1)
        plt.plot(t, mean_path, color='cyan', linewidth=2, label='Mean Path');
        ax.set(title="Monte Carlo simulated price paths\n$S_0$ = {}, $sigma$ = {}, $a$ = {}, $b$ = {}, $lambda$ = {}, Nsim = {}"\
               .format(S, sigma, Nsim), \
               xlabel='Time (days)', ylabel=f"{security}")
        plt.legend()
        plt.show()
    
    return mean_path

def pure_drift_diffusion(S, T, mu, sigma, Nsteps, Nsim, seed, plot=False):
    
    # Set random seed
    np.random.seed(seed)

    # Calculate the length of the time step
    Delta_t = T/Nsteps

    simulated_paths = np.zeros([Nsim, Nsteps+1])

    # Replace the first column of the array with the vector of initial price S
    simulated_paths[:,0] = S

    Z_1 = np.random.normal(size=[Nsim, Nsteps])

    # Populate the matrix with Nsim randomly generated paths of length Nsteps
    for i in range(Nsteps):
        simulated_paths[:,i+1] = simulated_paths[:,i]*np.exp((mu - sigma**2/2)*Delta_t +\
                                                             sigma*np.sqrt(Delta_t) * Z_1[:,i])

    mean_path = np.mean(simulated_paths, axis=0)
    
    if plot:
        sns.set(palette='rocket_r')
        plt.figure(figsize=(10,8))
        
        ax = plt.axes()
        t = np.linspace(0, T, Nsteps+1) * Nsteps
        jump_diffusion = ax.plot(t, simulated_paths.transpose());
        plt.setp(jump_diffusion, linewidth=1)
        plt.plot(t, mean_path, color='cyan', linewidth=2, label='Mean Path');
        ax.set(title="Monte Carlo simulated price paths\n$S_0$ = {}, $sigma$ = {}, $a$ = {}, $b$ = {}, $lambda$ = {}, Nsim = {}"\
               .format(S, sigma, Nsim), \
               xlabel='Time (days)', ylabel=f"{security}")
        plt.legend()
        plt.show()
    
    return mean_path

# In[6] Objective Functions
def model_1_objective_function(params, observed_prices, Nsim, seed):
    A_c0, A_d0, eta_prod, alpha, epsilon, Theta_good, c, sigma_carbon, Income, q_t, \
    sigma, Lambda, a, b = params
    
    carbon_return_array = relative_sector_evolution_params_only(A_c0, A_d0, eta_prod, alpha, epsilon, 
                                                    Theta_good, c, sigma_carbon, Income, q_t, \
                                                    len(observed_prices), plot=False)
    
    mean_path = economic_jump_diffusion(S=observed_prices[0], T=1, mu_array=carbon_return_array, 
                               sigma=sigma, Lambda=Lambda, a=a, b=b, Nsteps=len(observed_prices) - 1,\
                               Nsim=Nsim, seed=seed, plot=False)
    
    absolute_errors = np.abs((mean_path - observed_prices) / observed_prices)
    mape = np.mean(absolute_errors) * 100  # Convert to percentage

    return mape

def model_2_objective_function(params, observed_prices, Nsim, seed):
    mu, sigma, Lambda, a, b = params
    
    mean_path = pure_jump_diffusion(S=observed_prices[0], T=1, mu=mu, 
                               sigma=sigma, Lambda=Lambda, a=a, b=b, Nsteps=len(observed_prices) - 1,\
                               Nsim=Nsim, seed=seed, plot=False)
    
    absolute_errors = np.abs((mean_path - observed_prices) / observed_prices)
    mape = np.mean(absolute_errors) * 100  # Convert to percentage

    return mape

def model_3_objective_function(params, observed_prices, Nsim, seed):
    A_c0, A_d0, eta_prod, alpha, epsilon, Theta_good, c, sigma_carbon, Income, q_t, sigma = params
    
    carbon_return_array = relative_sector_evolution_params_only(A_c0, A_d0, eta_prod, alpha, epsilon, 
                                                    Theta_good, c, sigma_carbon, Income, q_t, \
                                                    len(observed_prices), plot=False)
    
    mean_path = economic_drift_diffusion(S=observed_prices[0], T=1, mu_array=carbon_return_array, 
                               sigma=sigma, Nsteps=len(observed_prices) - 1, Nsim=Nsim, seed=seed, plot=False)
    
    absolute_errors = np.abs((mean_path - observed_prices) / observed_prices)
    mape = np.mean(absolute_errors) * 100  # Convert to percentage

    return mape

def model_4_objective_function(params, observed_prices, Nsim, seed):
    mu, sigma = params
    
    mean_path = pure_drift_diffusion(S=observed_prices[0], T=1, mu=mu, 
                               sigma=sigma, Nsteps=len(observed_prices) - 1, Nsim=Nsim, seed=seed, plot=False)
    
    absolute_errors = np.abs((mean_path - observed_prices) / observed_prices)
    mape = np.mean(absolute_errors) * 100  # Convert to percentage

    return mape

# In[7] Optimization, Results, Plotting
def model_1_minimization_results_and_plotting(Nsim, seed, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    bounds = [(0.001, 10.0), (0.001, 10.0), (0.001, 1.0), (0.001, 1.0), (0.001, 3.0), 
              (1.001, 3.0), (0.001, 3.0), (0.001, 3.0), (1.0, 100.0), (0.001, 1.0),
              (0.1, 0.50), (0.000001, 0.09), (-0.50, 0.50), (-0.001, 0.50)]

    initial_guess = [1.0, 1.5, 0.2, 0.2, 2.0, 1.1, 0.10, 1.1, 10, 0.5, 0.5, 0.005, 0.1, 0.1]

    result = minimize(model_1_objective_function, initial_guess, args=(in_sample_data, Nsim_fixed, seed_fixed),\
                      method='Powell', bounds=bounds)

    optimized_params = result.x
    economic_params = optimized_params[:10]  # Economic model parameters
    simulation_params = optimized_params[10:]  # Simulation model parameters
    sigma_opt, Lambda_opt, a_opt, b_opt = simulation_params

    economic_evolution = relative_sector_evolution_params_only(*economic_params, num_periods=total_days, plot=False)

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = economic_jump_diffusion(S=credit_prices.iloc[0], T=1, mu_array=economic_evolution, 
                                        sigma=sigma_opt, Lambda=Lambda_opt, a=a_opt, b=b_opt, 
                                        Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100
    
    parameter_names = ["A_c0", "A_d0", "η", "alpha", "ε", "θ", 
                       "c", "σ_carbon", "I", "q", "σ", "λ", "µ_jump", "δ2"]

    optimized_params_with_names = [(param_name, round(param_value, 5)) for param_name, param_value in zip(parameter_names, optimized_params)]

    jump_model_params = ["σ", "λ", "µ_jump", "δ2"]

    
    if plot:
        # Plotting all simulated paths based on economic evolution
        plt.figure(figsize=(10, 6))
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
        
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            plt.plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
    
        # Plotting in-sample data, mean simulated path, and closest simulated path
        plt.plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        plt.plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        plt.plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        plt.axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
    
        jump_model_params_str = ', '.join([f"{param} = {optimized_params[10+i]:.5f}" for i, param in enumerate(jump_model_params)])
    
        plt.title(f'Model 1 Simulated Price Path \n{jump_model_params_str}')
    
        plt.xlabel('Time (Days)')
        plt.ylabel(f"{security}")
        plt.legend()
        plt.grid(True)
        plt.show()

    print("Model 1 MAPE for Mean Simulated Path:", f"{mape_mean_simulated:.3f}%")
    print("Model 1 MAPE for Mean Simulated Path In Out-of-Sample Only:", f"{mape_mean_simulated_out:.3f}%")
    print("Model 1 MAPE for Closest Simulated Path:", f"{min_MAPE:.3f}%")
    
    return optimized_params_with_names, mape_mean_simulated, mape_mean_simulated_out, min_MAPE

def model_2_minimization_results_and_plotting(Nsim, seed, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    bounds = [(-5.0, 5.0), (0.1, 0.50), (0.000001, 0.09), (-0.50, 0.50), (-0.001, 0.50)]

    initial_guess = [0.5, 0.5, 0.005, 0.1, 0.1]

    result = minimize(model_2_objective_function, initial_guess, args=(in_sample_data, Nsim_fixed, seed_fixed),\
                      method='Powell', bounds=bounds)

    optimized_params = result.x
    mu_opt, sigma_opt, Lambda_opt, a_opt, b_opt = optimized_params

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = pure_jump_diffusion(S=credit_prices.iloc[0], T=1, mu=mu_opt, 
                                        sigma=sigma_opt, Lambda=Lambda_opt, a=a_opt, b=b_opt, 
                                        Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100
    
    parameter_names = ["µ", "σ", "λ", "µ_jump", "δ2"]

    optimized_params_with_names = [(param_name, round(param_value, 5)) for param_name, param_value in zip(parameter_names, optimized_params)]

    jump_model_params = ["µ", "σ", "λ", "µ_jump", "δ2"]
    
    if plot:
        # Plotting all simulated paths based on economic evolution
        plt.figure(figsize=(10, 6))
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
    
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            plt.plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
    
        # Plotting in-sample data, mean simulated path, and closest simulated path
        plt.plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        plt.plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        plt.plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        plt.axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
    
        jump_model_params_str = ', '.join([f"{param} = {optimized_params[0+i]:.5f}" for i, param in enumerate(jump_model_params)])
    
        plt.title(f'model 2 Simulated Price Path \n{jump_model_params_str}')
    
        plt.xlabel('Time (Days)')
        plt.ylabel(f"{security}")
        plt.legend()
        plt.grid(True)
        plt.show()

    print("model 2 MAPE for Mean Simulated Path:", f"{mape_mean_simulated:.3f}%")
    print("model 2 MAPE for Mean Simulated Path In Out-of-Sample Only:", f"{mape_mean_simulated_out:.3f}%")
    print("model 2 MAPE for Closest Simulated Path:", f"{min_MAPE:.3f}%")
    
    return optimized_params_with_names, mape_mean_simulated, mape_mean_simulated_out, min_MAPE

def model_3_minimization_results_and_plotting(Nsim, seed, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    bounds = [(0.001, 10.0), (0.001, 10.0), (0.001, 1.0), (0.001, 1.0), (0.001, 3.0), 
              (1.001, 3.0), (0.001, 3.0), (0.001, 3.0), (1.0, 100.0), (0.001, 1.0), (0.1, 0.50)]

    initial_guess = [1.0, 1.5, 0.2, 0.2, 2.0, 1.1, 0.10, 1.1, 10, 0.5, 0.5]

    result = minimize(model_3_objective_function, initial_guess, args=(in_sample_data, Nsim_fixed, seed_fixed),\
                      method='Powell', bounds=bounds)

    optimized_params = result.x
    economic_params = optimized_params[:10]  # Economic model parameters
    simulation_params = optimized_params[10:]  # Simulation model parameters
    sigma_opt = simulation_params
  
    economic_evolution = relative_sector_evolution_params_only(*economic_params, num_periods=total_days, plot=False)

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = economic_drift_diffusion(S=credit_prices.iloc[0], T=1, mu_array=economic_evolution, 
                                        sigma=sigma_opt, Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100
    
    parameter_names = ["A_c0", "A_d0", "η", "alpha", "ε", "θ", 
                       "c", "σ_carbon", "I", "q", "σ"]

    optimized_params_with_names = [(param_name, round(param_value, 5)) for param_name, param_value in zip(parameter_names, optimized_params)]

    jump_model_params = ["σ"]
    
    if plot:
        # Plotting all simulated paths based on economic evolution
        plt.figure(figsize=(10, 6))
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
    
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            plt.plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
    
        # Plotting in-sample data, mean simulated path, and closest simulated path
        plt.plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        plt.plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        plt.plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        plt.axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
    
        jump_model_params_str = ', '.join([f"{param} = {optimized_params[10+i]:.5f}" for i, param in enumerate(jump_model_params)])
    
        plt.title(f'model 3 Simulated Price Path \n{jump_model_params_str}')
    
        plt.xlabel('Time (Days)')
        plt.ylabel(f"{security}")
        plt.legend()
        plt.grid(True)
        plt.show()

    print("model 3 MAPE for Mean Simulated Path:", f"{mape_mean_simulated:.3f}%")
    print("model 3 MAPE for Mean Simulated Path In Out-of-Sample Only:", f"{mape_mean_simulated_out:.3f}%")
    print("model 3 MAPE for Closest Simulated Path:", f"{min_MAPE:.3f}%")
    
    return optimized_params_with_names, mape_mean_simulated, mape_mean_simulated_out, min_MAPE

def model_4_minimization_results_and_plotting(Nsim, seed, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    bounds = [(-5.0, 5.0), (0.1, 0.50)]

    initial_guess = [0.5, 0.5]

    result = minimize(model_4_objective_function, initial_guess, args=(in_sample_data, Nsim_fixed, seed_fixed),\
                      method='Powell', bounds=bounds)

    optimized_params = result.x
    mu_opt, sigma_opt = optimized_params

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = pure_drift_diffusion(S=credit_prices.iloc[0], T=1, mu=mu_opt, 
                                        sigma=sigma_opt, Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100
    
    parameter_names = ["µ", "σ"]

    optimized_params_with_names = [(param_name, round(param_value, 5)) for param_name, param_value in zip(parameter_names, optimized_params)]

    jump_model_params = ["µ", "σ"]
    
    if plot:
        # Plotting all simulated paths based on economic evolution
        plt.figure(figsize=(10, 6))
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
    
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            plt.plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
    
        # Plotting in-sample data, mean simulated path, and closest simulated path
        plt.plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        plt.plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        plt.plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        plt.axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
    
        jump_model_params_str = ', '.join([f"{param} = {optimized_params[0+i]:.5f}" for i, param in enumerate(jump_model_params)])
    
        plt.title(f'model 4 Simulated Price Path \n{jump_model_params_str}')
    
        plt.xlabel('Time (Days)')
        plt.ylabel(f"{security}")
        plt.legend()
        plt.grid(True)
        plt.show()

    print("model 4 MAPE for Mean Simulated Path:", f"{mape_mean_simulated:.3f}%")
    print("model 4 MAPE for Mean Simulated Path In Out-of-Sample Only:", f"{mape_mean_simulated_out:.3f}%")
    print("model 4 MAPE for Closest Simulated Path:", f"{min_MAPE:.3f}%")
    
    return optimized_params_with_names, mape_mean_simulated, mape_mean_simulated_out, min_MAPE

# In[8] Running Simulations Nruns times Functions
def run_model_1(Nruns, Nsim_fixed, seed_fixed):
    model_para_results = []
    model_mean_mape_results = []
    model_mean_mape_out_results = []

    for _ in range(Nruns):
        Model_1_opt_params, Model_1_mape_mean, Model_1_mape_mean_out, Model_1_min_mape = \
            model_1_minimization_results_and_plotting(Nsim=Nsim_fixed, seed=seed_fixed, plot=False)
            
        model_para_results.append((Model_1_opt_params))
        model_mean_mape_results.append((Model_1_mape_mean))
        model_mean_mape_out_results.append((Model_1_mape_mean_out))
    
    avg_model_para_results = []
    for param_index in range(len(model_para_results[0])):
        param_name = model_para_results[0][param_index][0]
        param_values = [params[param_index][1] for params in model_para_results]
        avg_param_value = sum(param_values) / len(param_values)
        avg_model_para_results.append((param_name, avg_param_value))

    avg_model_mean_mape_results = sum(model_mean_mape_results) / Nruns
    avg_model_mean_mape_out_results = sum(model_mean_mape_out_results) / Nruns
    
    return avg_model_para_results, avg_model_mean_mape_results, avg_model_mean_mape_out_results

def run_model_2(Nruns, Nsim_fixed, seed_fixed):
    model_para_results = []
    model_mean_mape_results = []
    model_mean_mape_out_results = []

    for _ in range(Nruns):
        model_2_opt_params, model_2_mape_mean, model_2_mape_mean_out, model_2_min_mape = \
            model_2_minimization_results_and_plotting(Nsim=Nsim_fixed, seed=seed_fixed, plot=False)
            
        model_para_results.append((model_2_opt_params))
        model_mean_mape_results.append((model_2_mape_mean))
        model_mean_mape_out_results.append((model_2_mape_mean_out))
    
    avg_model_para_results = []
    for param_index in range(len(model_para_results[0])):
        param_name = model_para_results[0][param_index][0]
        param_values = [params[param_index][1] for params in model_para_results]
        avg_param_value = sum(param_values) / len(param_values)
        avg_model_para_results.append((param_name, avg_param_value))

    avg_model_mean_mape_results = sum(model_mean_mape_results) / Nruns
    avg_model_mean_mape_out_results = sum(model_mean_mape_out_results) / Nruns
    
    return avg_model_para_results, avg_model_mean_mape_results, avg_model_mean_mape_out_results

def run_model_3(Nruns, Nsim_fixed, seed_fixed):
    model_para_results = []
    model_mean_mape_results = []
    model_mean_mape_out_results = []

    for _ in range(Nruns):
        model_3_opt_params, model_3_mape_mean, model_3_mape_mean_out, model_3_min_mape = \
            model_3_minimization_results_and_plotting(Nsim=Nsim_fixed, seed=seed_fixed, plot=False)
            
        model_para_results.append((model_3_opt_params))
        model_mean_mape_results.append((model_3_mape_mean))
        model_mean_mape_out_results.append((model_3_mape_mean_out))
    
    avg_model_para_results = []
    for param_index in range(len(model_para_results[0])):
        param_name = model_para_results[0][param_index][0]
        param_values = [params[param_index][1] for params in model_para_results]
        avg_param_value = sum(param_values) / len(param_values)
        avg_model_para_results.append((param_name, avg_param_value))

    avg_model_mean_mape_results = sum(model_mean_mape_results) / Nruns
    avg_model_mean_mape_out_results = sum(model_mean_mape_out_results) / Nruns
    
    return avg_model_para_results, avg_model_mean_mape_results, avg_model_mean_mape_out_results

def run_model_4(Nruns, Nsim_fixed, seed_fixed):
    model_para_results = []
    model_mean_mape_results = []
    model_mean_mape_out_results = []

    for _ in range(Nruns):
        model_4_opt_params, model_4_mape_mean, model_4_mape_mean_out, model_4_min_mape = \
            model_4_minimization_results_and_plotting(Nsim=Nsim_fixed, seed=seed_fixed, plot=False)
            
        model_para_results.append((model_4_opt_params))
        model_mean_mape_results.append((model_4_mape_mean))
        model_mean_mape_out_results.append((model_4_mape_mean_out))
    
    avg_model_para_results = []
    for param_index in range(len(model_para_results[0])):
        param_name = model_para_results[0][param_index][0]
        param_values = [params[param_index][1] for params in model_para_results]
        avg_param_value = sum(param_values) / len(param_values)
        avg_model_para_results.append((param_name, avg_param_value))

    avg_model_mean_mape_results = sum(model_mean_mape_results) / Nruns
    avg_model_mean_mape_out_results = sum(model_mean_mape_out_results) / Nruns
    
    return avg_model_para_results, avg_model_mean_mape_results, avg_model_mean_mape_out_results

# In[9] Running Simulations Nruns times
avg_model_1_para_results, avg_model_1_mean_mape_results, avg_model_1_mean_mape_out_results =\
    run_model_1(Nruns, Nsim_fixed, seed_fixed)
avg_model_2_para_results, avg_model_2_mean_mape_results, avg_model_2_mean_mape_out_results =\
    run_model_2(Nruns, Nsim_fixed, seed_fixed)
avg_model_3_para_results, avg_model_3_mean_mape_results, avg_model_3_mean_mape_out_results =\
    run_model_3(Nruns, Nsim_fixed, seed_fixed)
avg_model_4_para_results, avg_model_4_mean_mape_results, avg_model_4_mean_mape_out_results =\
    run_model_4(Nruns, Nsim_fixed, seed_fixed)

# In[10] Final Plots Function
def model_1_avg_plotting(Nsim, seed, avg_model_1_para_results, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    # Extracting parameter values from avg_model_1_para_results
    optimized_params = [value for _, value in avg_model_1_para_results]
    economic_params = optimized_params[:10]  # Economic model parameters
    simulation_params = optimized_params[10:]  # Simulation model parameters
    sigma_opt, Lambda_opt, a_opt, b_opt = simulation_params

    economic_evolution = relative_sector_evolution_params_only(*economic_params, num_periods=total_days, plot=False)

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = economic_jump_diffusion(S=credit_prices.iloc[0], T=1, mu_array=economic_evolution, 
                                        sigma=sigma_opt, Lambda=Lambda_opt, a=a_opt, b=b_opt, 
                                        Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100
    
    jump_model_params = ["σ", "λ", "µ_jump", "δ2"]

    # Compute log returns for Monte Carlo simulation
    simulated_log_returns = np.diff(np.log(all_simulated_paths), axis=1)

    # Compute log returns for real credit data
    real_log_returns = np.diff(np.log(credit_prices.values))

    if plot:
        # Compute PDFs using KDE for Monte Carlo simulation
        simulated_pdf = gaussian_kde(simulated_log_returns.flatten())

        # Compute PDFs using KDE for real credit data
        real_pdf = gaussian_kde(real_log_returns.flatten())

        # Plotting both plots combined
        fig, axs = plt.subplots(2, 1, figsize=(17, 9), gridspec_kw={'height_ratios': [4, 1]})

        # Plotting the price plot
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            axs[0].plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
        axs[0].plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        axs[0].plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        axs[0].plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        axs[0].axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
        axs[0].set_xlabel('Time (Days)')
        axs[0].set_ylabel(f"{security}")
        axs[0].legend()
        axs[0].grid(True)

        # Plotting the distribution plot
        x = np.linspace(np.min(real_log_returns), np.max(real_log_returns), 1000)
        axs[1].plot(x, real_pdf(x), label='Real Credit Data', color='blue')
        axs[1].plot(x, simulated_pdf(x), label='Monte Carlo Simulation', color='red')
        axs[1].set_xlabel('Log Returns')
        axs[1].set_ylabel('Probability Density')
        axs[1].legend()
        axs[1].grid(True)

        plt.tight_layout()
        
        # Adding a title to the entire plot
        fig.suptitle('Model 1 Optimized Simulation', fontsize=20)
        
        plt.show()
    
    ks_statistic, p_value = kstest(simulated_pdf(x), real_pdf(x))
       
    return ks_statistic, p_value
    
def model_2_avg_plotting(Nsim, seed, avg_model_2_para_results, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    optimized_params = [value for _, value in avg_model_2_para_results]
    mu_opt, sigma_opt, Lambda_opt, a_opt, b_opt = optimized_params

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = pure_jump_diffusion(S=credit_prices.iloc[0], T=1, mu=mu_opt, 
                                        sigma=sigma_opt, Lambda=Lambda_opt, a=a_opt, b=b_opt, 
                                        Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100
    
    jump_model_params = ["µ", "σ", "λ", "µ_jump", "δ2"]
    # Compute log returns for Monte Carlo simulation
    simulated_log_returns = np.diff(np.log(all_simulated_paths), axis=1)

    # Compute log returns for real credit data
    real_log_returns = np.diff(np.log(credit_prices.values))

    if plot:
        # Compute PDFs using KDE for Monte Carlo simulation
        simulated_pdf = gaussian_kde(simulated_log_returns.flatten())

        # Compute PDFs using KDE for real credit data
        real_pdf = gaussian_kde(real_log_returns.flatten())

        # Plotting both plots combined
        fig, axs = plt.subplots(2, 1, figsize=(17, 9), gridspec_kw={'height_ratios': [4, 1]})

        # Plotting the price plot
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            axs[0].plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
        axs[0].plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        axs[0].plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        axs[0].plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        axs[0].axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
        axs[0].set_xlabel('Time (Days)')
        axs[0].set_ylabel(f"{security}")
        axs[0].legend()
        axs[0].grid(True)

        # Plotting the distribution plot
        x = np.linspace(np.min(real_log_returns), np.max(real_log_returns), 1000)
        axs[1].plot(x, real_pdf(x), label='Real Credit Data', color='blue')
        axs[1].plot(x, simulated_pdf(x), label='Monte Carlo Simulation', color='red')
        axs[1].set_xlabel('Log Returns')
        axs[1].set_ylabel('Probability Density')
        axs[1].legend()
        axs[1].grid(True)

        plt.tight_layout()
        
        # Adding a title to the entire plot
        fig.suptitle('model 2 Optimized Simulation', fontsize=20)
        
        plt.show()
    
    ks_statistic, p_value = kstest(simulated_pdf(x), real_pdf(x))
       
    return ks_statistic, p_value

def model_3_avg_plotting(Nsim, seed, avg_model_3_para_results, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    optimized_params = [value for _, value in avg_model_3_para_results]
    economic_params = optimized_params[:10]  # Economic model parameters
    simulation_params = optimized_params[10:]  # Simulation model parameters
    sigma_opt = simulation_params[0]
    
    economic_evolution = relative_sector_evolution_params_only(*economic_params, num_periods=total_days, plot=False)

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = economic_drift_diffusion(S=credit_prices.iloc[0], T=1, mu_array=economic_evolution, 
                                        sigma=sigma_opt, Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100
    
    jump_model_params = ["σ"]
    
    # Compute log returns for Monte Carlo simulation
    simulated_log_returns = np.diff(np.log(all_simulated_paths), axis=1)

    # Compute log returns for real credit data
    real_log_returns = np.diff(np.log(credit_prices.values))

    if plot:
        # Compute PDFs using KDE for Monte Carlo simulation
        simulated_pdf = gaussian_kde(simulated_log_returns.flatten())

        # Compute PDFs using KDE for real credit data
        real_pdf = gaussian_kde(real_log_returns.flatten())

        # Plotting both plots combined
        fig, axs = plt.subplots(2, 1, figsize=(17, 9), gridspec_kw={'height_ratios': [4, 1]})

        # Plotting the price plot
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            axs[0].plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
        axs[0].plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        axs[0].plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        axs[0].plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        axs[0].axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
        axs[0].set_xlabel('Time (Days)')
        axs[0].set_ylabel(f"{security}")
        axs[0].legend()
        axs[0].grid(True)

        # Plotting the distribution plot
        x = np.linspace(np.min(real_log_returns), np.max(real_log_returns), 1000)
        axs[1].plot(x, real_pdf(x), label='Real Credit Data', color='blue')
        axs[1].plot(x, simulated_pdf(x), label='Monte Carlo Simulation', color='red')
        axs[1].set_xlabel('Log Returns')
        axs[1].set_ylabel('Probability Density')
        axs[1].legend()
        axs[1].grid(True)

        plt.tight_layout()
        
        # Adding a title to the entire plot
        fig.suptitle('model 3 Optimized Simulation', fontsize=20)
        
        plt.show()
    
    ks_statistic, p_value = kstest(simulated_pdf(x), real_pdf(x))
       
    return ks_statistic, p_value

def model_4_avg_plotting(Nsim, seed, avg_model_4_para_results, plot=False):
    Nsim_fixed = Nsim
    seed_fixed = seed

    optimized_params = [value for _, value in avg_model_4_para_results]
    mu_opt, sigma_opt = optimized_params  # Extracting mu and sigma from optimized parameters

    all_simulated_paths = []
    for i in range(Nsim_fixed):
        simulated_path = pure_drift_diffusion(S=credit_prices.iloc[0], T=1, mu=mu_opt, 
                                        sigma=sigma_opt, Nsteps=total_days-1, Nsim=1, seed=seed_fixed, plot=False)
        all_simulated_paths.append(simulated_path)
        
    mean_simulated_path = np.mean(all_simulated_paths, axis=0)

    min_MAPE = float('inf')
    closest_simulated_path = None
    for simulated_path in all_simulated_paths:
        mape = np.mean(np.abs((simulated_path - credit_prices.values) / credit_prices.values)) * 100
        if mape < min_MAPE:
            min_MAPE = mape
            closest_simulated_path = simulated_path

    # Calculate the MAPE for the mean simulated path
    mape_mean_simulated = np.mean(np.abs((mean_simulated_path - credit_prices.values) / credit_prices.values)) * 100

    mape_mean_simulated_out = np.mean(np.abs((mean_simulated_path[in_sample_length:] - credit_prices[in_sample_length:].values) / credit_prices[in_sample_length:].values)) * 100

    jump_model_params = ["µ", "σ"]
    
    # Compute log returns for Monte Carlo simulation
    simulated_log_returns = np.diff(np.log(all_simulated_paths), axis=1)

    # Compute log returns for real credit data
    real_log_returns = np.diff(np.log(credit_prices.values))

    if plot:
        # Compute PDFs using KDE for Monte Carlo simulation
        simulated_pdf = gaussian_kde(simulated_log_returns.flatten())

        # Compute PDFs using KDE for real credit data
        real_pdf = gaussian_kde(real_log_returns.flatten())

        # Plotting both plots combined
        fig, axs = plt.subplots(2, 1, figsize=(17, 9), gridspec_kw={'height_ratios': [4, 1]})

        # Plotting the price plot
        sns.set_palette("rocket_r")
        palette = sns.color_palette()
        for i in range(Nsim_fixed):
            simulated_path = all_simulated_paths[i]
            axs[0].plot(credit_prices.index, simulated_path, color=palette[i % len(palette)], alpha=0.1)
        axs[0].plot(credit_prices.index, credit_prices, color='black', linewidth=2, label=f"{security}")
        axs[0].plot(credit_prices.index, mean_simulated_path, color='Cyan', linewidth=2, label='Mean Simulated Path')
        axs[0].plot(credit_prices.index, closest_simulated_path, color='springgreen', linewidth=2, label='Closest Simulated Path')
        axs[0].axvline(x=credit_prices.index[in_sample_length], color='Purple', linestyle='--', label='Out-of-Sample Begins')
        axs[0].set_xlabel('Time (Days)')
        axs[0].set_ylabel(f"{security}")
        axs[0].legend()
        axs[0].grid(True)

        # Plotting the distribution plot
        x = np.linspace(np.min(real_log_returns), np.max(real_log_returns), 1000)
        axs[1].plot(x, real_pdf(x), label='Real Credit Data', color='blue')
        axs[1].plot(x, simulated_pdf(x), label='Monte Carlo Simulation', color='red')
        axs[1].set_xlabel('Log Returns')
        axs[1].set_ylabel('Probability Density')
        axs[1].legend()
        axs[1].grid(True)

        plt.tight_layout()
        
        # Adding a title to the entire plot
        fig.suptitle('model 4 Optimized Simulation', fontsize=20)
        
        plt.show()
    
    ks_statistic, p_value = kstest(simulated_pdf(x), real_pdf(x))
       
    return ks_statistic, p_value

# In[11] Final Plots
model_1_ks_stat, model_1_ks_stat_pval = model_1_avg_plotting(Nsim_fixed, seed_fixed, avg_model_1_para_results, plot=True)
model_2_ks_stat, model_2_ks_stat_pval = model_2_avg_plotting(Nsim_fixed, seed_fixed, avg_model_2_para_results, plot=True)
model_3_ks_stat, model_3_ks_stat_pval = model_3_avg_plotting(Nsim_fixed, seed_fixed, avg_model_3_para_results, plot=True)
model_4_ks_stat, model_4_ks_stat_pval = model_4_avg_plotting(Nsim_fixed, seed_fixed, avg_model_4_para_results, plot=True)

print(f"{model_1_ks_stat:.5f}", f"{model_1_ks_stat_pval:.5f}")
print(f"{model_2_ks_stat:.5f}", f"{model_2_ks_stat_pval:.5f}")
print(f"{model_3_ks_stat:.5f}", f"{model_3_ks_stat_pval:.5f}")
print(f"{model_4_ks_stat:.5f}", f"{model_4_ks_stat_pval:.5f}")

# In[12] Report
def run_report(file_name):
    document = Document()

    avg_model_1_para_results.insert(10, ("μ", "-"))

    models = [avg_model_1_para_results, avg_model_2_para_results,
              avg_model_3_para_results, avg_model_4_para_results]

    parameters_order = [param for param, _ in avg_model_1_para_results]

    mean_path_mapes = [avg_model_1_mean_mape_results, avg_model_2_mean_mape_results,
                       avg_model_3_mean_mape_results, avg_model_4_mean_mape_results]

    mean_path_mapes_out = [avg_model_1_mean_mape_out_results, avg_model_2_mean_mape_out_results,
                           avg_model_3_mean_mape_out_results, avg_model_4_mean_mape_out_results]

    table = document.add_table(rows=1, cols=len(models) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'

    for i, model_params in enumerate(models, start=1):
        hdr_cells[i].text = f'Model {i}'

    for param in parameters_order:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        for i, model_params in enumerate(models, start=1):
            value = next((value for param_name, value in model_params if param_name == param), '-')
            row_cells[i].text = str(f"{value:.3f}" if isinstance(value, float) else value)

    table.add_row().cells[0].text = 'Mean Path MAPE'
    for i, mape_mean in enumerate(mean_path_mapes):
        cell = table.rows[-1].cells[i + 1]
        cell.text = str(f"{mape_mean:.2f}%")
        if mape_mean == min(mean_path_mapes):
            cell.paragraphs[0].runs[0].bold = True

    table.add_row().cells[0].text = 'Out-Sample Mean Path MAPE'
    for i, mape_mean_out in enumerate(mean_path_mapes_out):
        cell = table.rows[-1].cells[i + 1]
        cell.text = str(f"{mape_mean_out:.2f}%")
        if mape_mean_out == min(mean_path_mapes_out):
            cell.paragraphs[0].runs[0].bold = True
    
    table.rows[11].cells[2].text = str(f"{avg_model_2_para_results[0][1]:.3f}")
    table.rows[11].cells[4].text = str(f"{avg_model_4_para_results[0][1]:.3f}")
    
    table.add_row().cells[0].text = 'KS Statistic'
    ks_stats = [model_1_ks_stat, model_2_ks_stat, model_3_ks_stat, model_4_ks_stat]
    ks_stat_pvals = [model_1_ks_stat_pval, model_2_ks_stat_pval, model_3_ks_stat_pval, model_4_ks_stat_pval]
    min_ks_stat = min(ks_stats)

    for i, (ks_stat, ks_stat_pval) in enumerate(zip(ks_stats, ks_stat_pvals)):
        cell = table.rows[-1].cells[i + 1]
        cell.text = str(f"{ks_stat:.3f}")
        # Add stars if p-value is less than 1%
        if ks_stat_pval < 0.01:
            cell.text += ' ***'
        if ks_stat == min_ks_stat:
            cell.paragraphs[0].runs[0].bold = True

    document.save(f"{file_name}.docx")
    avg_model_1_para_results.pop(10)
    return avg_model_1_para_results

run_report("Simulation Results, YBBZ4")
# In[13] End